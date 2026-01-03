import os

file_service_path = os.path.join("backend", "app", "services", "file_service.py")

# This new code effectively 'Scrubs' the data clean before generating PDF
robust_file_service = """import pystache
from weasyprint import HTML, CSS
import io
import docx
import re

# 1. HELPER: Clean pure Hex (Strip hashes, handle nulls)
def get_clean_hex(value, default="000000"):
    if not value:
        return default
    # Remove #, spaces, and quotes
    clean = str(value).strip().replace("#", "").replace('"', '').replace("'", "")
    # Check length
    if len(clean) not in [3, 6, 8]:
        return default # Invalid hex, revert to black to prevent crash
    return clean

def create_pdf_from_template(template_html: str, template_css: str, cv_data: dict) -> bytes:
    # 2. SANITIZE DATA
    # We strip hashes from input to avoid double-hash issues later
    # We guarantee these are safe hex strings "2c3e50"
    safe_accent = get_clean_hex(cv_data.get("accentColor") or cv_data.get("accent_color"), "2c3e50")
    safe_text = get_clean_hex(cv_data.get("textColor") or cv_data.get("text_color"), "333333")
    
    safe_font = cv_data.get("fontFamily") or cv_data.get("font_family") or "sans-serif"
    # Ensure fonts don't have dangerous characters that break CSS strings
    safe_font = safe_font.replace('"', '').replace(";", "")

    # 3. BUILD INJECTED CSS
    # We create a 'shim' CSS block that defines the variables properly with ONE hash
    css_variables = f\"\"\"
    :root {{
        --primary: #{safe_accent} !important;
        --text-color: #{safe_text} !important;
        --font: '{safe_font}' !important;
        --accent_color: #{safe_accent} !important; /* legacy support */
    }}
    body {{
        color: #{safe_text} !important;
        font-family: '{safe_font}', sans-serif !important;
    }}
    \"\"\"

    # 4. PREPARE TEMPLATE DATA
    # We pass the 'safe' versions too, but also pass the FULL hash version for
    # templates that might use direct insertion like {{full_accent_code}}
    render_data = {
        **cv_data,
        # Normalized safe values (No hash)
        "accent_color_clean": safe_accent,
        "text_color_clean": safe_text,
        # Standard values (With hash) - FOR USE IN HTML inline styles
        "accent_color": f"#{safe_accent}",
        "text_color": f"#{safe_text}",
        "font_family": safe_font,
        
        # HTML Helpers
        "experience": (cv_data.get("experience") or "").replace("\\n", "<br/>"),
        "education": (cv_data.get("education") or "").replace("\\n", "<br/>")
    }

    # 5. RENDER TEMPLATE STRINGS
    try:
        # Prepend our safe variable block to the template's CSS
        full_css_string = css_variables + "\\n" + template_css
        
        rendered_html = pystache.render(template_html, render_data)
        rendered_css = pystache.render(full_css_string, render_data)
        
        # 6. REGEX POLISH (The Final Safety Net)
        # Catch accidental ##123456 or # #123456 occurrences
        rendered_css = re.sub(r'#\s*#', '#', rendered_css)
        rendered_css = re.sub(r'#+([A-Fa-f0-9]{6})', r'#\1', rendered_css)
        
    except Exception as e:
        print(f"Template Rendering Error: {e}")
        # Return an error PDF page instead of crashing server 500
        error_html = f"<h1>PDF Gen Error</h1><p>{str(e)}</p>"
        return HTML(string=error_html).write_pdf()

    # 7. GENERATE PDF via WeasyPrint
    try:
        html = HTML(string=rendered_html)
        css = CSS(string=rendered_css)
        return html.write_pdf(stylesheets=[css], presentational_hints=True)
    except Exception as e:
        # Catch internal WeasyPrint parser errors
        print(f"WeasyPrint Syntax Error: {e}")
        # Often helpful to print the first few lines of CSS to debug
        print(f"CSS Context: {rendered_css[:100]}...")
        raise e

def create_docx_from_data(cv_data: dict) -> bytes:
    document = docx.Document()
    
    # Simple Title Header
    name = cv_data.get('fullName') or cv_data.get('full_name') or "Your Name"
    document.add_heading(name, 0)
    
    info = f"{cv_data.get('email', '')} | {cv_data.get('phone', '')} | {cv_data.get('jobTitle') or cv_data.get('job_title', '')}"
    document.add_paragraph(info)
    
    # Iterate standard sections
    sections = [
        ("Summary", cv_data.get('summary') or cv_data.get('professional_summary')),
        ("Experience", cv_data.get('experience') or cv_data.get('experience_points')),
        ("Education", cv_data.get('education') or cv_data.get('education_formatted')),
        ("Skills", cv_data.get('skills') or cv_data.get('suggested_skills'))
    ]
    
    for title, content in sections:
        if content:
            document.add_heading(title, 1)
            # Content cleaning
            if isinstance(content, list):
                for item in content: document.add_paragraph(item, style='List Bullet')
            else:
                clean = str(content).replace('<br/>', '\\n').replace('<ul>', '').replace('</ul>', '')
                clean = clean.replace('<li>', '• ').replace('</li>', '\\n')
                for line in clean.split('\\n'):
                    if line.strip(): document.add_paragraph(line.strip())

    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream.read()
"""

try:
    with open(file_service_path, "w", encoding="utf-8") as f:
        f.write(robust_file_service)
    print("✅ File Service REPLACED with Debugger Logic.")
    print("   - Includes 'get_clean_hex' helper.")
    print("   - Injects Safe CSS Variables.")
    print("   - Performs final Regex Cleanup on CSS.")
except Exception as e:
    print(f"❌ Error: {e}")