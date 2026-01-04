import pystache
from weasyprint import HTML, CSS
import io
import docx
import re

def ensure_clean_hex(color_value, default_hex="333333"):
    """
    Returns a clean 6-character HEX string WITHOUT the # prefix.
    This is for data that will be rendered into templates that have #{{color}}.
    
    Examples:
        "#2c3e50" -> "2c3e50"
        "2c3e50"  -> "2c3e50"
        "invalid" -> "333333" (default)
    """
    if not color_value:
        return default_hex

    s = str(color_value).strip()
    clean_hex = s.lstrip("#")

    if re.fullmatch(r"[0-9a-fA-F]{3}|[0-9a-fA-F]{6}", clean_hex):
        if len(clean_hex) == 3:
            clean_hex = ''.join([c*2 for c in clean_hex])
        return clean_hex

    return default_hex

def create_pdf_from_template(template_html: str, template_css: str, cv_data: dict) -> bytes:
    """
    Renders CV template with Mustache and generates PDF using WeasyPrint.
    
    Templates should have CSS like: --primary: #{{accent_color}};
    And we pass accent_color WITHOUT the # (e.g., "2c3e50").
    """
    
    # Extract and Clean Colors (NO # prefix - templates add it)
    accent = ensure_clean_hex(
        cv_data.get('accentColor') or cv_data.get('accent_color'), 
        "2c3e50"
    )
    text_col = ensure_clean_hex(
        cv_data.get('textColor') or cv_data.get('text_color'), 
        "333333"
    )
    
    # Extract and Sanitize Font
    font_name = cv_data.get('fontFamily') or cv_data.get('font_family') or "sans-serif"
    font_name = re.sub(r'[;"\'#]+', '', font_name).strip()

    # Prepare Render Data for Mustache
    render_data = {
        **cv_data,
        "accent_color": accent,
        "text_color": text_col,
        "font_family": font_name,
        "experience": (cv_data.get('experience') or '').replace('\n', '<br/>'),
        "education": (cv_data.get('education') or '').replace('\n', '<br/>'),
        "summary": (cv_data.get('summary') or '').replace('\n', '<br/>'),
        "skills": cv_data.get('skills') if isinstance(cv_data.get('skills'), list) else []
    }

    if isinstance(cv_data.get('skills'), str):
        render_data["skills"] = [s.strip() for s in cv_data.get('skills').split(',') if s.strip()]

    try:
        # Render Templates
        compiled_html = pystache.render(template_html, render_data)
        compiled_css = pystache.render(template_css, render_data)
        
        # CSS Post-Processing (Safety Net)
        if "##" in compiled_css:
            compiled_css = compiled_css.replace("##", "#")
        
        compiled_css = re.sub(r':\s*#\s*;', f': #{text_col};', compiled_css)
        compiled_css = re.sub(r':\s*#\s*\}', f': #{text_col}', compiled_css)
        compiled_css = re.sub(r'#\s+([0-9a-fA-F])', r'#\1', compiled_css)

        # Generate PDF
        doc = HTML(string=compiled_html)
        style = CSS(string=compiled_css)
        pdf_bytes = doc.write_pdf(stylesheets=[style], presentational_hints=True)
        
        return pdf_bytes

    except Exception as e:
        # Generate error PDF with diagnostic info
        err_html = f"""
        <html>
        <head>
            <style>
                body {{
                    font-family: monospace;
                    padding: 40px;
                    background-color: #fff5f5;
                    color: #c53030;
                }}
                h1 {{
                    border-bottom: 3px solid #c53030;
                    padding-bottom: 10px;
                }}
                .error-box {{
                    background: white;
                    border: 2px solid #fc8181;
                    border-radius: 8px;
                    padding: 20px;
                    margin: 20px 0;
                }}
                .debug-info {{
                    background: #f7fafc;
                    border-left: 4px solid #4299e1;
                    padding: 15px;
                    margin-top: 20px;
                    font-size: 12px;
                }}
            </style>
        </head>
        <body>
            <h1>⚠️ PDF Generation Error</h1>
            <div class="error-box">
                <strong>Error Message:</strong>
                <pre>{str(e)}</pre>
            </div>
            <div class="debug-info">
                <strong>Debug Information:</strong><br/>
                • Accent Color: {accent}<br/>
                • Text Color: {text_col}<br/>
                • Font Family: {font_name}<br/>
                <br/>
                <strong>Common Fixes:</strong><br/>
                1. Check that CSS template uses #{{{{accent_color}}}} (with hash)<br/>
                2. Verify color values are valid 6-char hex codes<br/>
                3. Run /api/setup_production to resync templates
            </div>
        </body>
        </html>
        """
        try:
            return HTML(string=err_html).write_pdf()
        except:
            return b"PDF Generation Failed - See Server Logs"


def create_docx_from_data(cv_data: dict) -> bytes:
    """
    Creates a Word document from CV data.
    """
    document = docx.Document()
    
    document.add_heading(cv_data.get('fullName', 'Candidate'), 0)
    document.add_paragraph(f"{cv_data.get('email','')} | {cv_data.get('phone','')}")
    
    sections = ['summary', 'experience', 'education', 'skills']
    for sec in sections:
        val = cv_data.get(sec, "")
        if val:
            document.add_heading(sec.capitalize(), 1)
            
            if isinstance(val, list):
                val = ", ".join(val)
            
            clean = str(val).replace("<br/>", "\n").replace("<ul>","").replace("</ul>","")
            clean = clean.replace("<li>", "• ").replace("</li>", "\n")
            clean = re.sub(r'<[^>]+>', '', clean)
            
            for line in clean.split('\n'):
                if line.strip(): 
                    document.add_paragraph(line.strip())

    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream.read()